import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import os

doc = docx.Document()

def add_heading_centered(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def add_bold_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.add_run(text).bold = True
    return p

# --- TITLE PAGE ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A\nProject Report\nOn\n")
run.font.size = Pt(14)
run.bold = True
run2 = p.add_run("AurSikho – Online Course Enrollment System\n\n")
run2.font.size = Pt(20)
run2.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Submitted by\n").bold = True
p.add_run("Ashish Baviskar (SRN: 31251941)\n")
p.add_run("Himanshi Ladda (SRN: 31251597)\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Div: B\nMCA – I     SEM – II\n\n").bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Under the guidance of\n").bold = True
p.add_run("Prof. Khushbu Kushwaha\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("For the Academic Year 2025-26\n\n").bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Vishwakarma University\nKondhwa, Pune.").bold = True

doc.add_page_break()

# --- CERTIFICATE ---
add_heading_centered('CERTIFICATE', level=1)
p = doc.add_paragraph()
p.add_run('This is to certify that our team has successfully completed the project work entitled "AurSikho – Online Course Enrollment System" in partial fulfilment of MCA – I SEM – II for the year 2025-2026. The team has worked under my guidance and direction. This project represents their original work and efforts.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\nDate: 23 April 2026\nPlace: Pune')
doc.add_paragraph('\n\nGuide Signature: ___________________\n(Prof. Khushbu Kushwaha)')
doc.add_page_break()

# --- DECLARATION ---
add_heading_centered('DECLARATION', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.add_run('We certify that the work contained in this report is original and has been done by us under the guidance of our guide. The work has not been submitted to any other Institute for any degree or diploma. We have followed the guidelines provided by the Institute in preparing the report. We have conformed to the norms and guidelines given in the Ethical Code of Conduct of the Institute.')

doc.add_paragraph('\nName and Signature of Project Team Members:\n\n1. Ashish Baviskar (SRN: 31251941) ______________\n\n2. Himanshi Ladda (SRN: 31251597) ______________')
doc.add_page_break()

# --- INDEX ---
add_heading_centered('INDEX', level=1)
index_items = [
    "CHAPTER 1: INTRODUCTION",
    "   1.1 Abstract",
    "   1.2 Existing System and Need for System",
    "   1.3 Scope of System",
    "   1.4 Operating Environment",
    "   1.5 Brief Description of Technology Used",
    "CHAPTER 2: PROPOSED SYSTEM",
    "   2.1 Feasibility Study",
    "   2.2 Objectives of the Proposed System",
    "CHAPTER 3: ANALYSIS AND DESIGN",
    "   3.1 Entity Relationship Diagram (ERD)",
    "   3.2 Use Case Diagram",
    "   3.3 Data Flow Diagram (DFD)",
    "CHAPTER 4: IMPLEMENTATION",
    "   4.1 Modules Description",
    "CHAPTER 5: CONCLUSION"
]
for item in index_items:
    doc.add_paragraph(item)
doc.add_page_break()

# --- CHAPTER 1 ---
doc.add_heading('CHAPTER 1: INTRODUCTION', level=1)
doc.add_heading('1.1 Abstract', level=2)
doc.add_paragraph(
    '"AurSikho" is an innovative online course enrollment system designed to facilitate digital learning in a completely free and accessible manner. '
    'It provides a comprehensive platform where instructors can easily create, manage, and distribute courses. The platform natively supports hybrid '
    'lesson formats, seamlessly integrating YouTube video embeds with rich text descriptions, as well as complex multiple-choice quizzes. '
    'Students can browse a diverse catalog, filter courses by category or difficulty, enroll, and track their progress lesson by lesson. '
    'A key feature of AurSikho is its automated credentialing system: upon successful completion of all lessons and achieving a passing grade on the final quiz, '
    'students are automatically issued a visually appealing, dynamically generated PNG digital certificate. This certificate includes a unique UUID, allowing anyone '
    'to verify its authenticity via a public verification portal. Ultimately, the platform aims to bridge the gap between expert knowledge and eager learners '
    'through a highly responsive, glassmorphic web interface built on a robust Python backend.'
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('1.2 Existing System and Need for System', level=2)
doc.add_paragraph(
    'Existing traditional education systems and early-generation LMS (Learning Management System) platforms often suffer from being overly complex, highly expensive, '
    'and lacking in automated, verifiable credentialing. First, high costs for premium courses limit accessibility for many students globally. Second, '
    'manual tracking of student progress is extremely tedious for instructors, leading to administrative bottlenecks. Furthermore, the lack of automated, verifiable '
    'certificate generation means institutions must spend significant time verifying student credentials manually via email or phone calls.\n\n'
    'Therefore, there is a pressing need for a streamlined, free-to-use platform that entirely automates progress tracking and certificate issuance. '
    'A system that provides an engaging hybrid learning experience while removing the financial and administrative barriers to entry is crucial for modern educational paradigms.'
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('1.3 Scope of System', level=2)
scope_items = [
    'Role-Based Access Control: Secure and distinct dashboards and permissions for Students, Instructors, and Administrators.',
    'Advanced Course Management: Instructors can create Draft and Published courses, assign metadata (category, difficulty), and monitor student enrollment.',
    'Hybrid Lesson Content: Native support for lessons consisting of Video embeds, Markdown/Rich Text, or a hybrid of both simultaneously.',
    'Automated Dynamic Assessments: Creation of MCQs with customizable passing thresholds and automatic grading.',
    'Real-time Progress Tracking: Visual progress bars indicating completion percentage and specific lessons passed.',
    'Verifiable Certificates: Automated generation of personalized PNG certificates with a unique Verification UID.',
    'Responsive User Interface: Mobile-first design principles using Bootstrap 5 with modern gradients and animations.'
]
for item in scope_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.4 Operating Environment', level=2)
doc.add_paragraph('Server-side Requirement: Standard Cloud Server (e.g., Render, Railway) or Local Machine, Python 3.9+, SQLite3 Database (upgradable to PostgreSQL).')
doc.add_paragraph('Client-side Requirement: Any modern web browser (Google Chrome, Safari, Mozilla Firefox) with an active Internet connection. No dedicated app installation is required.')

doc.add_heading('1.5 Brief Description of Technology Used', level=2)
tech_items = [
    'Python 3.9+: The primary programming language used for all backend business logic, route handlers, and mathematical operations.',
    'Flask Framework: A lightweight WSGI web application framework used to build the backend logic, handle URL routing, and manage application state.',
    'SQLite3 & Flask-SQLAlchemy: A relational database paired with a robust Object Relational Mapper (ORM) to handle queries without writing raw SQL.',
    'Bootstrap 5 & Custom CSS: A powerful front-end framework used to create responsive, mobile-first web pages with custom CSS defining glassmorphic aesthetics.',
    'Pillow (PIL): The Python Imaging Library used for programmatically drawing customized completion certificates (text overlays, font adjustments).',
    'Jinja2: A fast, expressive, extensible templating engine utilized to render dynamic HTML pages on the server side.',
    'Mermaid JS: Used for generating architectural and flow diagrams.'
]
for item in tech_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# --- CHAPTER 2 ---
doc.add_heading('CHAPTER 2: PROPOSED SYSTEM', level=1)
doc.add_heading('2.1 Feasibility Study', level=2)
doc.add_paragraph(
    'Technical Feasibility: The project utilizes established, open-source technologies (Python, Flask, SQLite, Bootstrap) that are well-documented and robust. '
    'The required hardware is minimal, ensuring smooth development and deployment. The skills required to maintain the system are standard web development skills.\n\n'
    'Economic Feasibility: Built entirely on open-source software, the system requires absolutely zero software licensing fees. It can be hosted on free-tier '
    'or extremely low-cost cloud platforms, making it highly cost-effective compared to commercial enterprise LMS solutions.\n\n'
    'Operational Feasibility: The intuitive, role-segregated user interface ensures that students can easily learn and instructors can easily teach without requiring '
    'advanced technical skills or prolonged training periods.'
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('2.2 Objectives of the Proposed System', level=2)
obj_items = [
    'To provide a completely free, globally accessible digital learning environment.',
    'To empower instructors with simple tools to create hybrid (video + text) lessons effortlessly.',
    'To fully automate student progress tracking, reducing instructor overhead.',
    'To accurately evaluate students via dynamic MCQ quizzes before granting completion status.',
    'To programmatically issue visually appealing, professional certificates of completion.',
    'To allow public, cryptographic-style verification of issued certificates to prevent fraud.'
]
for item in obj_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# --- CHAPTER 3 ---
doc.add_heading('CHAPTER 3: ANALYSIS AND DESIGN', level=1)
doc.add_paragraph('This section details the architectural design of the AurSikho platform, illustrating data models, user interactions, and information flow.')

doc.add_heading('3.1 Entity Relationship Diagram (ERD)', level=2)
doc.add_paragraph('The ERD illustrates the relationships between core entities: Users (Students/Instructors), Courses, Lessons, Quizzes, Enrollments, and Certificates. It highlights the relational integrity of the database.')
try:
    if os.path.exists('ProjectReport/erd.png'):
        doc.add_picture('ProjectReport/erd.png', width=Inches(6.0))
    else:
        doc.add_paragraph("[ERD diagram image generating...]")
except Exception as e:
    doc.add_paragraph(f"[Could not load ERD diagram: {e}]")

doc.add_heading('3.2 Use Case Diagram', level=2)
doc.add_paragraph('The Use Case Diagram defines the interactions between the three primary actors (Student, Instructor, Admin) and the system modules, ensuring access control restrictions are visually mapped.')
try:
    if os.path.exists('ProjectReport/usecase.png'):
        doc.add_picture('ProjectReport/usecase.png', width=Inches(6.0))
    else:
        doc.add_paragraph("[Use Case diagram image generating...]")
except Exception as e:
    doc.add_paragraph(f"[Could not load Use Case diagram: {e}]")

doc.add_heading('3.3 Data Flow Diagram (DFD)', level=2)
doc.add_paragraph('The Data Flow Diagram (Level 0) represents the high-level movement of data between external entities (Users) and internal processes (Authentication, Database, Certificate Generator).')
try:
    if os.path.exists('ProjectReport/dfd.png'):
        doc.add_picture('ProjectReport/dfd.png', width=Inches(5.0))
    else:
        doc.add_paragraph("[DFD diagram image generating...]")
except Exception as e:
    doc.add_paragraph(f"[Could not load DFD diagram: {e}]")

doc.add_page_break()

# --- CHAPTER 4 ---
doc.add_heading('CHAPTER 4: IMPLEMENTATION', level=1)
doc.add_heading('4.1 Modules Description', level=2)
doc.add_paragraph('The application is divided into five core functional modules based on the App Factory pattern using Flask Blueprints:')

add_bold_paragraph('1. Authentication & Authorization Module')
doc.add_paragraph('Manages user registration, login, session handling, and role verification. Passwords are securely hashed using the PBKDF2 algorithm via Werkzeug. Custom Python decorators protect routes ensuring students cannot access instructor panels and vice versa.')

add_bold_paragraph('2. Course Management Module')
doc.add_paragraph('Provides a comprehensive dashboard for instructors to define courses, set difficulty metrics and categories, and manage state transitions (Draft vs. Published). It calculates dynamic statistics such as total enrolled students and total content duration.')

add_bold_paragraph('3. Lesson & Content Engine Module')
doc.add_paragraph('Supports the creation of Video, Text, and Hybrid lessons. It incorporates logic to automatically parse standard YouTube watch URLs and convert them into iframe-ready embed URLs, abstracting technical details from the instructor.')

add_bold_paragraph('4. Assessment & Progress Module')
doc.add_paragraph('Handles the creation of multiple-choice questions. It provides real-time progress tracking bars for students. When a quiz is submitted, it auto-grades the answers against the database and calculates pass/fail metrics based on the instructor-defined threshold.')

add_bold_paragraph('5. Dynamic Certificate Module')
doc.add_paragraph('Integrates with the Pillow (PIL) library to dynamically draw student names, course titles, and dates onto a master PNG template using programmatic coordinate mapping and font sizing. It generates a unique UUID (v4) for every certificate, saving it to the database for future public verification lookups.')

doc.add_page_break()

# --- CHAPTER 5 ---
doc.add_heading('CHAPTER 5: CONCLUSION', level=1)
doc.add_paragraph(
    'The "AurSikho" Online Course Enrollment System provides a highly robust, scalable, and user-friendly platform for free digital education. '
    'By automating critical, time-consuming tasks such as progress tracking, assessments, and certificate generation, it effectively eliminates '
    'administrative overhead for instructors and creates a seamless, rewarding learning journey for students. \n\n'
    'The use of modern web technologies—specifically the Python Flask framework paired with a custom Bootstrap 5 design system—ensures a highly performant '
    'and visually engaging experience across all devices, from desktop to mobile. As digital education continues to evolve, platforms like AurSikho '
    'demonstrate the feasibility of delivering premium learning experiences and verifiable credentials entirely free of charge.'
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('ProjectReport/Project_Report_Aursikho.docx')
print("Detailed Docx created successfully!")
