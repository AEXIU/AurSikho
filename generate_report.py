import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# Title Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("A\nProject Report\nOn\n").bold = True
p.add_run("AurSikho – Online Course Enrollment System\n\n").bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Submitted by\n").bold = True
p.add_run("Ashish Baviskar (SRN: 31251941)\n")
p.add_run("Himanshi Ladda (SRN: 31251597)\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Div: B\nMCA – I     SEM – II\n\n").bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Under the guidance of\n").bold = True
p.add_run("Prof. Khushbu Kushwaha\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("For the Academic Year 2025-26\n\n").bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Vishwakarma University\nKondhwa, Pune.").bold = True

doc.add_page_break()

# Certificate
doc.add_heading('CERTIFICATE', level=1)
doc.add_paragraph('This is to certify that our team has successfully completed the project work entitled "AurSikho – Online Course Enrollment System" in partial fulfilment of MCA – I SEM – II for the year 2025-2026. The team has worked under my guidance and direction.')
doc.add_paragraph('Date: 23 April 2026\nPlace: Pune')
doc.add_paragraph('\n\nGuide Signature: ___________________\n(Prof. Khushbu Kushwaha)')

doc.add_page_break()

# Declaration
doc.add_heading('DECLARATION', level=1)
doc.add_paragraph('We certify that the work contained in this report is original and has been done by us under the guidance of our guide.')
doc.add_paragraph('The work has not been submitted to any other Institute for any degree or diploma.')
doc.add_paragraph('We have followed the guidelines provided by the Institute in preparing the report and conformed to the norms and guidelines given in the Ethical Code of Conduct of the Institute.')
doc.add_paragraph('\nName and Signature of Project Team Members:\n1. Ashish Baviskar (SRN: 31251941) ______________\n2. Himanshi Ladda (SRN: 31251597) ______________')

doc.add_page_break()

# Chapters
doc.add_heading('CHAPTER 1: INTRODUCTION', level=1)
doc.add_heading('1.1 Abstract', level=2)
doc.add_paragraph('"AurSikho" is an online course enrollment system designed to facilitate digital learning. It provides a platform where instructors can create and manage courses, including hybrid video and text lessons, as well as multiple-choice quizzes. Students can browse, enroll in, and complete these courses completely free of charge. Upon successful completion of all lessons and the final quiz, students are automatically issued a verifiable digital certificate.')

doc.add_heading('1.2 Existing System and Need for System', level=2)
doc.add_paragraph('Existing traditional education systems and early-generation LMS platforms often suffer from being overly complex, expensive, and lacking in automated credentialing. There is a need for a streamlined, free-to-use platform that automates progress tracking and certificate issuance while providing a rich, hybrid learning experience.')

doc.add_heading('1.3 Scope of System', level=2)
doc.add_paragraph('Role-Based Access for Students, Instructors, and Administrators.\nCourse Management for hybrid content (Video + Text).\nAutomated Assessments via quizzes.\nVerifiable Certificates dynamically generated.')

doc.add_heading('1.4 Operating Environment', level=2)
doc.add_paragraph('Server-side: Python 3.9+, SQLite3 Database.\nClient-side: Any modern web browser.')

doc.add_heading('1.5 Brief Description of Technology Used', level=2)
doc.add_paragraph('Python Flask: Web framework.\nSQLite3 & SQLAlchemy: Database & ORM.\nBootstrap 5: Responsive frontend.\nPillow: Certificate image generation.')

doc.add_heading('CHAPTER 2: PROPOSED SYSTEM', level=1)
doc.add_heading('2.1 Feasibility Study', level=2)
doc.add_paragraph('Technical: Uses open-source tools.\nEconomic: Zero licensing costs.\nOperational: Intuitive interface.')

doc.add_heading('2.2 Objectives of the Proposed System', level=2)
doc.add_paragraph('Provide free digital learning.\nAutomate student tracking and assessments.\nIssue and verify completion certificates.')

doc.add_heading('CHAPTER 3: ANALYSIS AND DESIGN', level=1)
doc.add_paragraph('Please refer to the Markdown file (Project_Report_Aursikho.md) for the rendered UML Mermaid diagrams (ERD, Use Case, and DFD).')

doc.add_heading('CHAPTER 4: IMPLEMENTATION', level=1)
doc.add_heading('4.1 Modules', level=2)
doc.add_paragraph('1. Authentication Module\n2. Course Management Module\n3. Lesson & Content Module\n4. Assessment Module\n5. Certificate Module')

doc.add_heading('CHAPTER 5: CONCLUSION', level=1)
doc.add_paragraph('AurSikho provides a robust, scalable, and user-friendly platform for free digital education, eliminating administrative overhead through automated credentialing.')

doc.save('ProjectReport/Project_Report_Aursikho.docx')
print("Docx created!")
